import datetime
import time

import data
from core.SlackAPIClient import Slack
from logsetup import configure_log
from docx import Document
from core.AWSClient import S3

REPEAT_FREQUENCY = 10 * 60 * 1000
logger = configure_log(__name__)
slack = Slack()
s3 = S3()

date_time_format = lambda d: datetime.datetime.strftime(d, "%y%m%d")


class MeetingsHandler(object):
    """
    Handles the scheduling, ending, creation and update of docs for ended meetings
    """
    def start_meetings(self):
        meetings = data.find_meetings_to_schedule()
        if meetings:
            for meeting in meetings:
                self.invite_to_channel(meeting)
                slack.slack_client.api_call(
                    "chat.postMessage",
                    channel=meeting.slack_channel_id,
                    text=self.get_meeting_start_message(meeting),
                    as_user=False,
                    username="Meeting Scheduler"
                )

    def invite_to_channel(self, meeting):
        # Find Channel Members and Check if any of the users are yet to join the channel
        channel_members = slack.slack_client.api_call(
            "conversations.members",
            channel=meeting.slack_channel_id,
        ).get("members")
        users = []
        for guest in meeting.guests:
            if guest.user.slack_user_id in channel_members:
                logger.info(guest.user.slack_user_id +" already added to channel "+meeting.slack_channel_id)
                continue
            users.append(guest.user.slack_user_id)

        # Invite the users who are yet to join
        if len(users) > 0:
            channel_invite = slack.slack_client.api_call(
                "conversations.invite",
                channel=meeting.slack_channel_id,
                users=users
            )
            error = slack.check_error(channel_invite)
            if error:
                logger.error \
                    ("Failed to Invite Guests to Channel " + meeting.slack_channel_name + ". Error " + str(error))

    def get_meeting_start_message(self, meeting):
        message = []
        message.append(meeting.title)
        message.append(" meeting will be started at ")
        message.append(str(meeting.start_date.hour))
        message.append(":")
        message.append(str(meeting.start_date.minute))
        message.append(". Please find the meeting docs at ")
        message.append(meeting.meeting_docs_parent_folder)
        return ''.join(message)

    def end_meetings(self):
        meetings = data.find_meetings_to_end()
        if meetings:
            for meeting in meetings:
                slack.slack_client.api_call(
                    "chat.postMessage",
                    channel=meeting.slack_channel_id,
                    text=self.get_meeting_end_message(meeting),
                    as_user=False,
                    username="Meeting Scheduler"
                )

    def get_meeting_end_message(self, meeting):
        message = []
        message.append("Thanks for your participation in the meeting ")
        message.append(meeting.title)
        message.append(". Please find the meeting docs at ")
        message.append(meeting.meeting_docs_parent_folder)
        return ''.join(message)

    def update_meeting_docs(self):
        meeting_instances = data.find_meeting_instance_for_docs_update()
        # TODO: Handle when new users are manually added to channel
        users = data.get_user_map()
        if meeting_instances:
            for meeting_instance in meeting_instances:
                logger.info("Docs Update pending for the meeting "+meeting_instance.title)
                meeting_minutes = []
                has_more = True
                cursor = ''
                while has_more:
                    conversations = slack.slack_client.api_call(
                        "conversations.history",
                        channel=meeting_instance.slack_channel_id,
                        inclusive = True,
                        oldest = meeting_instance.start_time.timestamp(),
                        latest = meeting_instance.end_time.timestamp(),
                        cursor = cursor,
                    )
                    self.process_conversations(meeting_minutes,conversations,users)
                    has_more = conversations.get('has_more')
                    metadata = conversations.get('response_metadata')
                    if metadata:
                        cursor = metadata.get('next_cursor')
                self.create_meeting_doc(meeting_instance,meeting_minutes)
                data.update_meeting_instance_docs(meeting_instance)
                logger.info(meeting_instance.title+ " pending for Docs Update")
        else:
            logger.info("No Meetings pending for Docs Update")

    def process_conversations(self,meeting_minutes,conversations,users):
        for message in conversations.get("messages"):
            message_subtype = message.get("subtype")
            if not message_subtype:
                # TODO: Handle Mentions
                user_message = users.get(message.get("user"))+":"+message.get("text")
                files = message.get("files")
                if files:
                    for file in files:
                        meeting_minutes.append(file.get("permalink"))
                meeting_minutes.append(user_message)
            # elif message_subtype and message_subtype == 'bot_message':
            #     meeting_minutes.append(message.get("text"))

    def create_meeting_doc(self,meeting_instance,meeting_minutes):
        document = Document()
        title = meeting_instance.title + " - " + date_time_format(meeting_instance.start_time)
        document.add_heading(title)
        while len(meeting_minutes) > 0:
            document.add_paragraph(meeting_minutes.pop())
        meeting_minutes_doc = "./../meeting_minutes/" + title + ".docx"
        document.save(meeting_minutes_doc)
        upload_location = meeting_instance.meeting_docs_folder
        upload_location = upload_location[upload_location.index("buckets/")+len("buckets/"):]
        upload_bucket = upload_location[0:upload_location.index("/")]
        upload_key = upload_location[upload_location.index("/") + 1:]+"/"+title+".docx"
        data = open(meeting_minutes_doc, 'rb')
        s3.s3_client.Bucket(upload_bucket).put_object(Key=upload_key, Body=data)
        data.close()

if __name__ == '__main__':
    meeting_scheduler = MeetingsHandler()
    # We will be running this in an infinite loop no matter what environment as CronJob is still an alpha feature
    while True:
        meeting_scheduler.start_meetings()
        meeting_scheduler.end_meetings()
        meeting_scheduler.update_meeting_docs()
        time.sleep(REPEAT_FREQUENCY)
